from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def cover(document: Document, section: dict):
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(section.get("title", "Title"))
    run.bold = True
    run.font.size = Pt(26)

    subtitle_text = section.get("subtitle")
    if subtitle_text:
        document.add_paragraph()
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(subtitle_text)
        run.font.size = Pt(16)
        run.font.color.rgb = None

    document.add_page_break()


def toc(document: Document, section: dict):
    heading = document.add_heading(section.get("title", "Table of Contents"), level=1)
    toc_para = document.add_paragraph(
        "Table of contents will be inserted here."
    )
    document.add_page_break()


def body(document: Document, section: dict):
    heading = document.add_heading(section.get("title", "Section"), level=1)
    content = section.get("content", "")
    paragraphs = content.split("\n")
    for para_text in paragraphs:
        if para_text.strip():
            document.add_paragraph(para_text.strip())


def headers_footers(document: Document, section: dict):
    for section_obj in document.sections:
        header = section_obj.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.text = section.get("header_text", "")

        footer = section_obj.footer
        footer.is_linked_to_previous = False
        pf = footer.paragraphs[0]
        pf.text = section.get("footer_text", "")
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
